# backend/doc_utils.py
from __future__ import annotations

import io
import os
from typing import List, Tuple, Optional

# PDF
import fitz  # PyMuPDF

# DOCX
try:
    from docx import Document  # python-docx
except Exception:
    Document = None

# OCR (optional)
try:
    import pytesseract
    from PIL import Image
except Exception:
    pytesseract = None
    Image = None


def _clean_lines(lines: List[str]) -> str:
    out: List[str] = []
    for ln in lines:
        s = (ln or "").strip()
        if not s:
            continue
        # collapse internal whitespace a bit
        s = " ".join(s.split())
        if s:
            out.append(s)
    return "\n".join(out)


def extract_pdf_text(
    data: bytes,
    use_ocr: bool = False,
    ocr_langs: str = "eng",
) -> Tuple[str, int, int, int]:
    """
    Returns: (full_text, total_pages, ocr_pages, empty_pages)
    """
    ocr_pages = 0
    empty_pages = 0
    total_pages = 0
    parts: List[str] = []

    with fitz.open(stream=data, filetype="pdf") as doc:
        total_pages = doc.page_count
        for i in range(total_pages):
            page = doc.load_page(i)
            txt = (page.get_text() or "").strip()

            if not txt and use_ocr and pytesseract and Image:
                # rasterize and OCR
                pix = page.get_pixmap(dpi=200, alpha=False)
                img = Image.open(io.BytesIO(pix.tobytes("png")))
                txt = pytesseract.image_to_string(img, lang=ocr_langs) or ""
                if txt.strip():
                    ocr_pages += 1

            if not txt.strip():
                empty_pages += 1
            parts.append(txt or "")

    return _clean_lines(parts), total_pages, ocr_pages, empty_pages


def _docx_walk(document) -> List[str]:
    """
    Extract paragraphs + table cell text from DOCX.
    """
    out: List[str] = []

    # normal paragraphs
    for p in document.paragraphs:
        t = (p.text or "").strip()
        if t:
            out.append(t)

    # tables
    for tbl in document.tables:
        for row in tbl.rows:
            for cell in row.cells:
                t = (cell.text or "").strip()
                if t:
                    out.append(t)

    # headers/footers (best effort)
    for s in getattr(document, "sections", []):
        for header in (getattr(s, "header", None), getattr(s, "footer", None)):
            if not header:
                continue
            for p in getattr(header, "paragraphs", []):
                t = (p.text or "").strip()
                if t:
                    out.append(t)

    return out


def extract_docx_text(data: bytes) -> str:
    if not Document:
        return ""
    # python-docx needs a real file-like object; BytesIO works
    f = io.BytesIO(data)
    doc = Document(f)
    lines = _docx_walk(doc)
    return _clean_lines(lines)


def extract_text_from_bytes(
    filename: str,
    data: bytes,
    use_ocr: bool = False,
    ocr_langs: str = "eng",
) -> Tuple[str, int, int, int]:
    """
    Unified API.
    Returns: (full_text, pages, ocr_pages, empty_pages)
    For DOCX pages/ocr/empty are 0.
    """
    name = (filename or "").lower()
    if name.endswith(".pdf"):
        return extract_pdf_text(data, use_ocr=use_ocr, ocr_langs=ocr_langs)

    if name.endswith(".docx"):
        text = extract_docx_text(data)
        return text, 0, 0, 0

    # try image OCR if asked
    if use_ocr and pytesseract and Image:
        try:
            img = Image.open(io.BytesIO(data))
            txt = pytesseract.image_to_string(img, lang=ocr_langs)
            return (txt or ""), 1, 1 if (txt or "").strip() else 0, 1 if not (txt or "").strip() else 0
        except Exception:
            pass

    # plain text fallback
    try:
        return data.decode("utf-8", errors="ignore"), 0, 0, 0
    except Exception:
        return "", 0, 0, 0


def simple_chunks(
    text: str,
    chunk_size: int = 1200,
    overlap: int = 120,
    min_chars: int = 80,
) -> List[str]:
    """
    Lightweight chunker:
    - split by blank lines / headings
    - then slide a window with overlap
    """
    text = (text or "").strip()
    if not text:
        return []

    # First pass: split around likely paragraph breaks
    raw_parts = []
    for block in text.replace("\r\n", "\n").split("\n\n"):
        s = block.strip()
        if not s:
            continue
        raw_parts.append(s)

    # Merge tiny parts so they pass min_chars threshold
    parts: List[str] = []
    buf: List[str] = []
    acc = 0
    for p in raw_parts:
        if acc + len(p) < min_chars:
            buf.append(p)
            acc += len(p)
            continue
        if buf:
            buf.append(p)
            parts.append("\n\n".join(buf))
            buf, acc = [], 0
        else:
            parts.append(p)
    if buf:
        parts.append("\n\n".join(buf))

    # Second pass: slide a window to approx. chunk_size
    out: List[str] = []
    for p in parts:
        if len(p) <= chunk_size:
            out.append(p)
            continue
        start = 0
        while start < len(p):
            end = min(len(p), start + chunk_size)
            piece = p[start:end]
            if len(piece) >= min_chars:
                out.append(piece)
            if end == len(p):
                break
            start = max(end - overlap, start + 1)

    return out
